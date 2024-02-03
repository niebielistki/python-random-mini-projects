# doc-party-invite-creator.py
# Description: Generates custom invitations for each guest listed in a text file and saves them into a Word document.
# Each invitation is styled and placed on a separate page.

import docx
from docx.enum.text import WD_BREAK

# Initialize a new Word document
doc = docx.Document()

# Replace with the path to your guest list text file
guests_file_path = ""  # Replace with your actual file path

# Read guest names from the file and add invitations to the document
with open(guests_file_path, 'r') as file:
    for line in file:
        name = line.strip()  # Remove any leading/trailing whitespace

        # Add invitation text with appropriate styles
        doc.add_paragraph('It is my pleasure to invite you to my party', style='Heading 2')
        doc.add_paragraph(name, style='Heading 1')
        doc.add_paragraph("You are cordially invited to our event!", style='Normal')
        doc.add_paragraph("Date: 21.03.2023", style='Normal')

        # Insert a page break after each invitation, except the last one
        doc.paragraphs[-1].runs[0].add_break(WD_BREAK.PAGE)

# Save the document with invitations
doc.save('party_invitations.docx')
print("Invitations have been created and saved to 'party_invitations.docx'.")
